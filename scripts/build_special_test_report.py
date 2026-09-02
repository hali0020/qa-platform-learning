from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "专项测试实现与指标数据汇总.docx"

NAVY = "20364B"
BLUE = "2E74B5"
PALE = "E8EEF5"
LIGHT = "F5F7F9"
GRAY = "5B6570"
GREEN = "2F7D5C"
AMBER = "A56616"
RED = "A43D3D"
BLACK = "111111"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent=120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[i]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, size=10.5, bold=None, color=BLACK, italic=None, name="Microsoft YaHei") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_text(doc, text, *, size=10.5, bold=False, color=BLACK, align=None, before=0, after=6, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_with_next = keep
    if align is not None:
        p.alignment = align
    set_font(p.add_run(text), size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    set_font(p.add_run(text), size=10.2)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    set_font(p.add_run(text), size=10.2)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run(text), size={1: 16, 2: 13, 3: 11.5}[level], bold=True,
             color=BLUE if level < 3 else NAVY)
    return p


def add_table(doc, headers, rows, widths, font_size=8.7):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for i, h in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], PALE)
        p = table.rows[0].cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(h), size=font_size, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if len(table.rows) % 2 == 1:
                set_cell_shading(cells[i], LIGHT)
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            set_font(p.add_run(str(value)), size=font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, label, text, color=BLUE):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    t.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    set_table_geometry(t, [9360])
    set_cell_shading(t.cell(0, 0), "EEF4FA")
    p = t.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run(label + "："), size=10, bold=True, color=color)
    set_font(p.add_run(text), size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def setup_styles(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.78)
    sec.bottom_margin = Inches(0.72)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.3)
    sec.footer_distance = Inches(0.3)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, BLUE),
        ("Heading 2", 13, 14, 7, BLUE),
        ("Heading 3", 11.5, 10, 5, NAVY),
    ):
        s = doc.styles[name]
        s.font.name = "Microsoft YaHei"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Bullet 2", "List Number"):
        s = doc.styles[name]
        s.font.name = "Microsoft YaHei"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        s.font.size = Pt(10.2)


def add_header_footer(doc):
    for sec in doc.sections:
        h = sec.header.paragraphs[0]
        h.clear()
        h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_font(h.add_run("QA Platform Learning  |  专项测试汇总"), size=8.5, color=GRAY)
        f = sec.footer.paragraphs[0]
        f.clear()
        f.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(f.add_run("内部技术资料  ·  2026-09-02"), size=8, color=GRAY)


def build():
    doc = Document()
    setup_styles(doc)

    add_text(doc, "专项测试", size=13, bold=True, color=AMBER, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    add_text(doc, "实现与指标数据汇总", size=28, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_text(doc, "QA Platform Learning · 当前仓库实现、验证结果与专项方案边界",
             size=12.5, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=28)
    add_table(doc, ["版本", "统计日期", "证据范围"],
              [["v1.0", "2026-09-02", "当前工作区源码、自动化测试、现有专项手册与技术文档"]],
              [1500, 1800, 6060], font_size=9.2)
    add_callout(doc, "结论", "当前仓库已形成一套可运行的 QA 平台专项验证体系；本次后端全量回归 557 项全部通过。弱网注入、图形自动判定、游戏性能长稳和 Lua 生成属于本项目后续设计范围，当前仓库没有可声明为实测成绩的原始数据。", GREEN)
    add_text(doc, "阅读口径", size=11, bold=True, color=NAVY, before=8, after=5)
    add_bullet(doc, "“已实现”表示当前仓库中存在对应代码与自动化测试。")
    add_bullet(doc, "“已设计”表示由本项目自行整理了测试方法、采集字段和参考门槛，但尚未形成可执行工具。")
    add_bullet(doc, "所有百分比、时延、FPS、温度等非本次运行值均明确标注为口径、SLO 示例或实验室参考门槛。")
    doc.add_page_break()

    add_heading(doc, "1. 汇总概览", 1)
    add_heading(doc, "1.1 本次可复核数据", 2)
    add_table(doc, ["项目", "结果", "口径/说明"], [
        ["后端自动化回归", "557 passed", "2026-09-02 使用 backend/.venv 执行 python -m pytest -q"],
        ["执行耗时", "100.78 秒", "pytest 报告总耗时；不等同于接口性能压测数据"],
        ["失败/错误", "0 / 0", "进程退出码 0"],
        ["告警", "58", "PytestCollectionWarning：业务类名以 Test 开头，不影响用例通过；建议后续清理"],
        ["静态测试函数命中", "403 处", "rg 扫描 def test_/async def test_；参数化后实际执行为 557 项"],
    ], [2100, 1800, 5460], font_size=9)

    add_heading(doc, "1.2 专项成熟度矩阵", 2)
    add_table(doc, ["专项", "状态", "当前实现", "指标/证据"], [
        ["API 与异步 HTTP", "已实现", "httpx Provider、分阶段超时、轮询、错误映射、取消与幂等边界", "自动化回归覆盖；Provider 时延/结果指标"],
        ["Webhook 与可靠投递", "已实现", "HMAC、时窗、sequence、水位线、Outbox、租约/CAS、退避、死信", "重复/冲突/stale/gap/恢复测试"],
        ["调度、Worker、设备租约", "已实现", "任务 claim、心跳、续租、取消、重试、死信、Cron、设备独占", "状态快照与并发不变量测试"],
        ["质量度量", "已实现", "汇总、趋势、套件覆盖 API 与前端报表", "通过率、自动化覆盖、执行触达、失败提单覆盖"],
        ["可观测性", "已实现", "健康探针、Request ID、JSON 日志、Prometheus 指标", "HTTP/任务/设备/Provider 指标"],
        ["数据导入/配置表", "部分实现", "CSV/XLSX 模板、预检、SHA-256、逐行错误、公式拒绝、提交/导出", "安全导入已测；Lua 确定性生成未实现"],
        ["身份与安全", "已实现/隔离", "Session、RBAC、CSRF、OIDC/Secret Store 边界、Webhook 安全", "越权/重放/敏感信息边界测试"],
        ["对象存储与制品", "已实现/隔离", "本地/S3 边界、Artifact 摘要、quarantine/restore", "摘要、路径、补偿、下载安全头测试"],
        ["弱网/断网/多人同步", "已设计", "Profile、注入流程、协议对账方法已整理", "无当前仓库实测样本"],
        ["兼容/图形回归", "已设计", "已整理设备矩阵、截图基线、mask、diff 与人工复核方案", "无自动 diff/pass-fail；无当前实测样本"],
        ["性能/热力/温度/长稳", "已设计", "已整理采样字段、聚合方法与参考门槛", "无当前原始采样"],
    ], [1450, 1150, 3650, 3110], font_size=7.7)

    add_heading(doc, "2. 已实现专项：实现方式与验证指标", 1)
    add_heading(doc, "2.1 API 与异步 HTTP 可靠性", 2)
    add_text(doc, "实现集中在 pipeline Provider、Learning CI、runtime service 与 webhook 路由。Provider 通过统一 trigger/get/cancel 契约接入；外部 HTTP 不置于长事务中，而是采用“写意图/claim → 事务外请求 → CAS 结算”的三段式。")
    for item in [
        "分离连接、读取、写入、连接池超时；将错误映射为稳定的领域结果。",
        "触发使用持久 trigger intent 与 Idempotency-Key；重试复用相同幂等键。",
        "202/异步运行通过轮询与签名 Webhook 收敛；终态禁止倒退。",
        "取消、超时、未知结果与 reconcile_required 分开表达，避免把不确定当失败或成功。",
    ]:
        add_bullet(doc, item)
    add_table(doc, ["验证点", "判定/数据"], [
        ["重复请求", "同 correlation/idempotency 输入重放不重复产生副作用；不同输入冲突"],
        ["状态收敛", "poll/Webhook 只能合法推进；旧 sequence 记 stale，不覆盖权威快照"],
        ["Provider 指标", "qa_provider_requests_total{provider,operation,outcome}；qa_provider_request_duration_seconds{provider,operation}"],
        ["接口容量口径", "目标 QPS ≈ 峰值活跃用户 × 单用户每秒请求数 × 1.5～2；当前未做容量实测"],
    ], [2400, 6960], font_size=8.8)

    add_heading(doc, "2.2 Webhook、Outbox 与故障恢复", 2)
    add_text(doc, "Webhook 对原始 body 先执行 16 KiB 上限，再按独立 Secret、固定五分钟时窗与 HMAC 校验。事件 ID、body SHA-256、sequence、occurred-at、Connection UUID 和 correlation ID 共同用于识别重放、冲突、缺口与时间回退。")
    add_table(doc, ["机制", "实现", "关键验证"], [
        ["接收幂等", "event_id + body digest", "同 ID 同内容为 duplicate；同 ID 不同内容为 conflict"],
        ["顺序保护", "单 Run 递增 sequence + watermark", "stale 不倒退；gap 标记对账；权威快照可清除标记"],
        ["主动投递", "持久 Outbox + 独立 Worker", "claim、租约、CAS、指数退避、死信、手工 retry"],
        ["任务唤醒", "事务 outbox + RabbitMQ 固定提示", "消息不携带任务 ID/参数/凭据；数据库 claim 仍是执行授权"],
    ], [1550, 3400, 4410], font_size=8.2)
    add_callout(doc, "可靠性边界", "该模型提供至少一次风格的可恢复处理，不承诺端到端“恰好一次”；幂等 Handler、权威数据库状态与对账机制共同消化重复提示。", AMBER)

    add_heading(doc, "2.3 Worker、Scheduler 与设备并发", 2)
    add_text(doc, "任务系统实现独立 Worker/Scheduler、数据库租约、心跳续租、取消、可重试失败、死信和设备独占。PostgreSQL 设计使用 SKIP LOCKED 与数据库时钟；SQLite 保留单机教学入口。")
    add_bullet(doc, "claim/finalize 均校验 owner、token hash、version 与租约有效性，失租处理者不得覆盖新 owner。")
    add_bullet(doc, "设备 lease 具备排他性；任务、设备状态分别以固定低基数枚举暴露。")
    add_bullet(doc, "真实多实例 PostgreSQL/RabbitMQ、Worker 强杀、Broker/数据库中断尚未在当前无 Docker 机器实测。")

    add_heading(doc, "2.4 数据导入、配置与制品完整性", 2)
    add_text(doc, "当前平台实现版本化 CSV/XLSX 模板、只读预检、文件 SHA-256、逐行错误、公式拒绝、用户确认提交与导出核对；Run Artifact 实现 pending→ready/failed→deleted 状态与摘要复核。")
    add_table(doc, ["检查项", "当前实现/标准"], [
        ["输入安全", "文件类型/大小、表头、行级字段、公式拒绝、hash 绑定预检与提交"],
        ["数据一致性", "逐行结果、唯一性/引用类校验按模块实现；审计记录提交事实"],
        ["Artifact", "保存后复核大小与 SHA-256；删除先 quarantine，元数据失败则 restore"],
        ["未完成", "完整数值 Schema、跨表规则、确定性 Excel/CSV→Lua、Lua 语法/加载与反向对账"],
    ], [2050, 7310], font_size=8.8)

    add_heading(doc, "2.5 身份、安全与敏感信息边界", 2)
    add_text(doc, "实现覆盖本地 Session、RBAC、CSRF、OIDC 隔离模式、Secret Store 接口、Provider URL 限制、Webhook 签名、上传/下载安全与日志脱敏。Jenkins、GitLab、BK-CI 仅使用 Mock 契约或自建隔离服务，默认关闭真实连接。")
    add_bullet(doc, "未授权写操作、触发人自批质量门禁、Webhook 重放绕过、任意文件访问、Secret 进入日志/指标均按零容忍设计。")
    add_bullet(doc, "Request 日志不记录 body、Cookie、Authorization、查询字符串或具体资源路径。")
    add_bullet(doc, "安全测试范围仅限授权的本地/隔离环境；当前项目不声明合规级审计能力。")

    add_heading(doc, "2.6 可观测性专项", 2)
    add_table(doc, ["信号", "实现", "注意事项"], [
        ["存活/就绪", "/health/live 不访问数据库；/health/ready 对受限 backend 执行 SELECT 1", "失败返回 503；不泄露路径、SQL 或凭据"],
        ["HTTP", "请求总数、延迟直方图、in-flight", "route 使用模板；/metrics 自身不计数"],
        ["业务", "任务状态、设备状态、Provider 请求结果和耗时", "标签限定固定枚举，拒绝 ID/URL/异常正文"],
        ["日志", "安全 Request ID + JSON 结构化日志", "关闭日志不影响 Request ID 与指标采集"],
    ], [1400, 4550, 3410], font_size=8.2)
    add_text(doc, "当前 Prometheus 指标清单：", size=10.3, bold=True, color=NAVY, after=3)
    for item in [
        "qa_http_requests_total{method,route,status_code}",
        "qa_http_request_duration_seconds{method,route}",
        "qa_http_requests_in_flight",
        "qa_automation_tasks{state}；qa_devices{state}",
        "qa_provider_requests_total{provider,operation,outcome}；qa_provider_request_duration_seconds{provider,operation}",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2.7 质量度量专项", 2)
    add_text(doc, "质量报表由后端统一计算，前端展示分子、分母和百分比；空分母显示“—”，不伪造 0%。趋势只使用已完成执行的最终结果。")
    add_table(doc, ["指标", "当前口径"], [
        ["自动化覆盖率", "当前 active 且 automated 用例数 / 当前 active 用例数"],
        ["执行触达率", "当前范围内已执行用例数 / 当前 active 用例数"],
        ["执行通过率", "passed / (passed + failed)；blocked、skipped、not_run 不进入分母"],
        ["失败提单覆盖", "有关联缺陷的 failed/blocked 结果对 / 全部 failed/blocked 结果对"],
        ["趋势", "仅完成执行；按日/周聚合通过、失败、阻塞等最终结果"],
    ], [2200, 7160], font_size=8.8)

    add_heading(doc, "3. 游戏专项测试方案与指标口径", 1)
    add_callout(doc, "状态说明", "本章是本项目自行整理的专项执行方法与参考门槛。当前仓库没有设备采样原始文件，不能将这些数字写成项目实测结果。", AMBER)

    add_heading(doc, "3.1 弱网、断网与多人同步", 2)
    add_table(doc, ["Profile", "RTT/延迟参考", "抖动", "丢包", "下行/上行参考"], [
        ["基线 Wi-Fi", "20～50 ms", "≤10 ms", "≤0.5%", "≥10 / ≥5 Mbps"],
        ["良好移动网", "50～100 ms", "10～30 ms", "≤1%", "5～20 / 1～5 Mbps"],
        ["弱网", "100～200 ms", "30～80 ms", "1%～3%", "1～5 Mbps / 256 Kbps～1 Mbps"],
        ["严重弱网", "200～500 ms", "80～150 ms", "3%～10%", "256 Kbps～1 Mbps / 128～512 Kbps"],
        ["断网/黑洞", "无响应", "—", "100% 或单向", "0"],
    ], [1450, 1800, 1400, 1450, 3260], font_size=8.1)
    add_text(doc, "实现建议：Linux 网关 tc netem、HTTP/TCP 使用 Toxiproxy 或自建代理、Windows 使用受控 WinDivert 类工具、移动设备统一经受控路由；普通 HTTP 代理不能完整模拟 UDP。每次注入必须用探测/抓包复核真实参数。")
    add_bullet(doc, "采集 request_id/session/room_id/sequence、发送/确认时间、RTT、重传、丢包、重连次数、检测时间、恢复时间与前后状态快照。")
    add_bullet(doc, "参考判定：关键会话重连 ≤5～10 秒、状态在 1～3 个快照周期内收敛；正式值由玩法和协议确定。")
    add_bullet(doc, "硬标准：无无限 loading；重试有界；不重复扣款/发奖/建角/订单；最终状态与权威服务端一致。")

    add_heading(doc, "3.2 兼容、安装升级与图形", 2)
    add_text(doc, "兼容矩阵按真实用户占比与风险权重选择 OS、芯片/GPU、内存/存储、分辨率/DPI/刷新率、图形 API、画质/HDR/超分及输入外设。每台设备须保存完整环境指纹。")
    add_bullet(doc, "P0 新装和主升级路径参考要求 100% 通过；文件 hash 不一致、存档/货币丢失、不可恢复半新半旧状态为 0。")
    add_bullet(doc, "空间边界至少覆盖 R-1 MB、R、R+20%；下载 10%、50%、99%、校验、解压、切换和迁移阶段分别中断。")
    add_bullet(doc, "图形硬标准：Crash、黑屏、严重伪影、关键对象缺失为 0；轻微视觉差异才使用 mask + SSIM/像素差并人工复核。")
    add_callout(doc, "当前实现边界", "当前仓库尚未实现设备截图采集和自动视觉回归。后续由本项目自行实现时，应补齐 baseline、设备指纹、mask、图片 hash、自动差异、阈值及人工复核；完成前不能声明已具备自动图像判定。", RED)

    add_heading(doc, "3.3 性能、场景热力、温度与稳定性", 2)
    add_table(doc, ["领域", "原始指标", "重点统计"], [
        ["流畅度", "FPS、帧时间、Jank/BigJank、Game/Render/RHI 线程", "平均、P95/P99、低帧占比、最长连续低帧"],
        ["CPU/GPU", "利用率、频率、显存、温度", "峰值、持续饱和、降频时刻、显存增长"],
        ["内存", "PSS、堆、原生堆、Swap", "峰值、增长斜率、退出场景后的回收率"],
        ["功耗", "电流、电压、电池温度、容量", "平均功率、单位时间耗电、温升"],
        ["加载", "启动/场景/资源阶段耗时", "P50/P95、最长阶段、相对基线回退"],
        ["网络/磁盘", "收发字节与包、读写量", "平均/峰值速率、累计流量、突刺"],
        ["渲染", "triangles、draw calls、mesh calls、粒子数", "热区最大值及与低帧时间相关性"],
    ], [1300, 4600, 3460], font_size=8.2)
    add_table(doc, ["目标", "平均 FPS 示例", "P95 帧时间", "P99 帧时间"], [
        ["60 FPS", "≥55", "≤20 ms", "≤33.3 ms"],
        ["30 FPS", "≥28", "≤40 ms", "≤66.7 ms"],
    ], [2340, 2340, 2340, 2340], font_size=9)
    add_bullet(doc, "参考门槛：PSS 相对基线增长不超过 5%～10% 且无线性上升；持续 FPS 相对冷态下降不超过 10%；新版本 FPS 回退不超过 5%，加载/内存回退不超过 10%。")
    add_bullet(doc, "CPU 持续 >80%、GPU 持续 >90% 是定位信号，不单独判失败；电池 42～45°C 仅为关注区间，优先使用 OS thermal state、降频和厂商规范。")
    add_bullet(doc, "发布阻断场景 Crash/ANR/OOM 必须为 0；Crash-free session ≥99.9% 只能在足够真实样本下作为线上 SLO。")
    add_text(doc, "本项目热力图采样设计：", size=10.3, bold=True, color=NAVY, after=3)
    add_bullet(doc, "字段：Location、Rotation、FPS、RHITriangles、RHIDrawPrimitiveCalls、MeshDrawCalls、NiagaraNumParticles、NiagaraNumMeshVerts。")
    add_bullet(doc, "同位置多朝向聚合最低 FPS、最大 triangles 与最大 RHI draw calls；全局观察 mesh draw calls、粒子数和粒子 mesh 顶点最大值。")
    add_callout(doc, "实现要求", "保留原始样本与 hash，增加 P1 FPS/P95 帧时间、每点样本数、缺失点检查、单位校验、设备信息和基线差异；可视化颜色只用于定位，不能替代按本项目设备档位和场景制定的门禁。", RED)

    add_heading(doc, "3.4 配置表与 Lua 专项", 2)
    add_text(doc, "数值表建议把业务数据与字段规则分离，以 schema_version 管理模板；检查类型、空值、最小/最大、唯一、概率和、外键、跨表不变量、精度与存储极值。")
    add_bullet(doc, "硬约束示例：主键重复/空 ID 为 0；min_value ≤ max_value；start_time < end_time；禁止的循环依赖为 0。")
    add_bullet(doc, "概率统一为 [0,1] 或 [0,100]，池内合计为 1 或 100；误差必须由领域明确。")
    add_bullet(doc, "完整流水线应为：输入/hash/模板检查 → Schema 与跨表校验 → 稳定排序 → 确定性 Lua → 隔离运行时加载 → 反向对账 → manifest/hash/diff → CI 审批。")
    add_callout(doc, "当前边界", "当前仓库已实现安全导入，但没有完整 Excel/CSV/JSON→Lua 生成器，也没有 Lua 加载测试与制品门禁；只能表述为“已实现导入、已设计生成链路”。", AMBER)

    add_heading(doc, "4. 指标字典与统一计算口径", 1)
    add_table(doc, ["指标", "公式/口径", "使用约束"], [
        ["通过率", "通过用例数 / 已执行用例数 × 100%", "平台正式口径进一步限定 passed/(passed+failed)"],
        ["缺陷逃逸率", "发布后缺陷 / 发布前后缺陷总数 × 100%", "需统一统计周期与严重级别"],
        ["回归变化率", "(新版本值 - 基线值) / 基线值 × 100%", "设备、场景、画质、构建必须可比"],
        ["平均 FPS", "有效 FPS 样本之和 / 样本数", "必须同时报告 P95/P99 帧时间与低帧占比"],
        ["帧预算", "1000 / 目标 FPS（ms）", "60 FPS≈16.67 ms；30 FPS≈33.33 ms"],
        ["内存增长斜率", "(结束 PSS - 开始 PSS) / 运行时间", "结合多轮回收与单调趋势判断泄漏"],
        ["吞吐率", "时间窗内成功请求数 / 时间窗长度", "同时报告错误率、并发、资源与下游耗时"],
        ["网络速率", "字节增量 × 8 / 时间增量", "不能用累计字节直接当瞬时速率"],
    ], [1700, 3900, 3760], font_size=8.2)
    add_text(doc, "分位数必须基于原始样本排序计算，不能用平均值替代。对有预热、随机和设备抖动的指标，同条件至少运行 3 次，报告中位数、P95/P99、最大值和失败样本；容量测试应另行设计样本量与置信区间。")

    add_heading(doc, "5. 发布门禁与建议后续动作", 1)
    add_heading(doc, "5.1 可先采用的实验室门禁", 2)
    add_bullet(doc, "关键用例 100% 通过；总回归通过率参考 ≥95%～98%；新增 blocker/critical 为 0。")
    add_bullet(doc, "自动化 flaky rate 参考 <1%～2%；不能用高总通过率掩盖关键用例失败。")
    add_bullet(doc, "性能回退不超过项目预算；配置校验全过；输入/输出/Artifact hash 一致。")
    add_bullet(doc, "所有门禁必须携带环境、样本量、统计周期、阈值来源与基线版本。")
    add_heading(doc, "5.2 优先补齐顺序", 2)
    for item in [
        "清理 58 条 pytest 收集告警，并在 CI 中固定保存 JUnit/XML 与运行环境元数据。",
        "在可用 Docker 的个人隔离环境实测 PostgreSQL/RabbitMQ、多实例 claim、强杀、租约过期、Broker/数据库中断恢复。",
        "落地受控弱网 Profile 与协议时间线采集，先覆盖登录、下载、结算与断网恢复。",
        "定义统一性能原始数据格式，补场景轨迹、分位数、样本数、缺失点和基线 diff。",
        "为图形报告增加设备指纹、baseline/mask/hash、SSIM/像素差、人工审批与自动结论字段。",
        "补完整数值 Schema 与确定性 Lua 生成、隔离加载、反向对账和 CI Artifact 门禁。",
    ]:
        add_number(doc, item)

    add_heading(doc, "附录 A：主要证据位置", 1)
    add_table(doc, ["证据", "仓库路径"], [
        ["专项测试总方法与门槛", "docs/GAME_QA_TESTING_INTERVIEW_GUIDE.md"],
        ["项目实现总览", "README.md"],
        ["可观测性与指标", "docs/OBSERVABILITY.md；backend/app/observability/"],
        ["CI 编排与一致性", "docs/PHASE6B_6C_ORCHESTRATION.md；backend/app/pipeline/；backend/app/runtime/"],
        ["质量指标实现", "backend/app/services/quality.py；backend/app/api/routes/quality.py；frontend/src/views/QualityReportsView.vue"],
        ["数据导入", "backend/app/data_transfer/；backend/app/api/routes/data_transfer.py"],
        ["自动化验证", "backend/tests/（本次执行 557 passed）"],
        ["部署与未完成边界", "infra/；docs/NEXT_DEVELOPMENT.md；docs/SECURITY.md"],
    ], [2500, 6860], font_size=8.5)

    add_heading(doc, "附录 B：真实性声明", 1)
    add_text(doc, "本文按“已实现、已设计”区分完成度。本次唯一新增实测运行数据是后端 pytest 结果（557 passed，100.78 秒，58 warnings）。性能 FPS、P95/P99、弱网参数、温度和发布阈值均为本项目整理的实验室示例或项目建议，不代表当前代码在真实设备、真实容量或生产拓扑中的验收成绩。")

    add_header_footer(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
